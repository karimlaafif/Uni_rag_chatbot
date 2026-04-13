import os
import re
import hashlib
import unicodedata
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF
from bs4 import BeautifulSoup
import aiohttp
import docx
from PIL import Image
import torch
from datetime import datetime
from langchain_core.documents import Document
from sqlalchemy import create_engine, text

from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from data_pipeline.semantic_chunker import SemanticChunker
# Singleton CLIP partagé avec chain.py — évite de charger ~600 MB deux fois
from shared_models import get_clip_model

import logging
logger = logging.getLogger(__name__)

class DataIngestionPipeline:
    def __init__(self, qdrant_manager=None):
        self.qdrant_manager = qdrant_manager

        self.text_splitter = SemanticChunker(
            model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
            breakpoint_percentile=70.0,   # cut at the 30% lowest-similarity transitions
            max_chunk_chars=1500,          # hard ceiling per chunk
            min_chunk_chars=200,           # discard/merge chunks smaller than this
            fallback_chunk_size=512,       # fallback splitter params (unchanged)
            fallback_overlap=64,
        )

        # Réutilise l'instance CLIP déjà chargée par RAGChatbot si le serveur
        # tourne, ou la charge pour la première fois si appelé en standalone.
        self.clip_model, self.clip_preprocess, self.tokenizer = get_clip_model()
        
    def _hash_content(self, content: str) -> str:
        return hashlib.sha256(content.encode('utf-8')).hexdigest()

    async def extract_webpage(self, url: str) -> str:
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                html = await response.text()
                soup = BeautifulSoup(html, "html.parser")
                return soup.get_text(separator="\n", strip=True)

    def extract_pdf(self, file_path: str) -> str:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text

    def extract_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    def extract_txt(self, file_path: str) -> str:
        # Tentative UTF-8 d'abord, fallback latin-1 pour les fichiers legacy
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    def extract_image(self, file_path: str) -> Dict[str, Any]:
        image = self.clip_preprocess(Image.open(file_path)).unsqueeze(0)
        with torch.no_grad():
            image_features = self.clip_model.encode_image(image)
        return {"image_vector": image_features.tolist()[0]}

    # ------------------------------------------------------------------
    # Artefacts PDF à supprimer : numéros de page, en-têtes/pieds répétitifs
    # Adapté aux documents UIZ (FR/AR/EN)
    # ------------------------------------------------------------------
    _PDF_ARTIFACTS = re.compile(
        r'^\s*('
        r'page\s+\d+\s*(\/\s*\d+)?'       # "Page 3" ou "Page 3/47"
        r'|\d+\s*\/\s*\d+'                 # "3/47" seul
        r'|©.*$'                           # mentions de copyright
        r'|université\s+ibn\s+zohr.*$'     # en-tête UIZ répété
        r'|ibn\s+zohr\s+university.*$'
        r'|جامعة\s+ابن\s+زهر.*$'          # en-tête en arabe
        r'|[-–—]{3,}'                      # lignes de séparation (---/–––)
        r')\s*$',
        flags=re.IGNORECASE | re.MULTILINE,
    )

    def clean_text(self, text: str) -> str:
       
        return "\n".join([line.strip() for line in text.split("\n") if line.strip()])

    def preprocess_text(self, text: str, source: str = "") -> str:
        
        if not text or not text.strip():
            return ""

        original_len = len(text)

        # ── Étape 1 : Normalisation Unicode NFC ──────────────────────────
        text = unicodedata.normalize("NFC", text)

        # ── Étape 2 : Réassemblage des mots coupés par césure PDF ────────
        # "étudi-\nant" → "étudiant"  /  "Ibn-\nZohr" → "Ibn-Zohr" (conservé)
        text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)

        # ── Étape 3 : Suppression des artefacts PDF ───────────────────────
        text = self._PDF_ARTIFACTS.sub("", text)

        # ── Étape 4 : Normalisation des espaces ───────────────────────────
        # Remplace tabs et espaces multiples/insécables par un espace simple.
        # On travaille ligne par ligne pour ne pas toucher aux sauts de ligne.
        lines = text.split("\n")
        lines = [re.sub(r'[ \t\u00a0\u202f\u2009]+', ' ', line).strip() for line in lines]
        text = "\n".join(lines)

        # ── Étape 5 : Max 1 ligne vide consécutive ────────────────────────
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Supprimer aussi les lignes qui ne contiennent que des espaces
        text = "\n".join(line for line in text.split("\n") if line.strip() or line == "")
        text = text.strip()

        # ── Étape 6 : Déduplication des blocs répétitifs ─────────────────
        # On découpe en paragraphes (séparés par ligne vide) et on déduplique
        # en préservant l'ordre d'apparition.
        paragraphs = text.split("\n\n")
        seen: set = set()
        unique_paragraphs = []
        duplicates_removed = 0
        for para in paragraphs:
            key = re.sub(r'\s+', ' ', para.strip().lower())
            if key and key not in seen:
                seen.add(key)
                unique_paragraphs.append(para)
            elif key in seen:
                duplicates_removed += 1
        text = "\n\n".join(unique_paragraphs)

        final_len = len(text)
        reduction_pct = round((1 - final_len / original_len) * 100, 1) if original_len > 0 else 0
        logger.debug(
            f"[preprocess] {source or 'doc'} : "
            f"{original_len} → {final_len} chars "
            f"({reduction_pct}% réduit, {duplicates_removed} blocs dupliqués supprimés)"
        )

        return text

    def log_error(self, doc_id: str, step: str, message: str):
        import json
        error_log = {
            "doc_id": doc_id,
            "etape": step,
            "message": message,
            "timestamp": datetime.utcnow().isoformat()
        }
        with open("ingestion_errors.log", "a", encoding="utf-8") as f:
            f.write(json.dumps(error_log) + "\n")

    def process_document(
        self,
        content: str,
        metadata: Dict[str, Any],
        force_reindex: bool = False,
    ) -> List[Document]:
        try:
            content = self.preprocess_text(content, source=metadata.get("source", ""))
            content_hash = self._hash_content(content)

            # Skip documents already in the vector store (delta-sync).
            # force_reindex=True (mode full) contourne cette vérification.
            if (
                not force_reindex
                and self.qdrant_manager
                and self.qdrant_manager.hash_exists(content_hash)
            ):
                return []

            metadata['content_hash'] = content_hash
            metadata['timestamp'] = datetime.utcnow().isoformat()

            # SemanticChunker.create_documents already:
            #   - splits on meaning boundaries
            #   - adds chunk_index, chunk_total, chunk_preview to metadata
            #   - filters chunks smaller than min_chunk_chars internally
            docs = self.text_splitter.create_documents(content, metadata=metadata)

            return docs
        except Exception as e:
            self.log_error(metadata.get('source', 'unknown_source'), "process_document", str(e))
            return []
        
    def ingest_file(
        self,
        file_path: str,
        metadata: Dict[str, Any],
        force_reindex: bool = False,
    ) -> List[Document]:
        try:
            ext = file_path.split('.')[-1].lower()
            if ext == 'pdf':
                content = self.extract_pdf(file_path)
                return self.process_document(content, metadata, force_reindex)
            elif ext == 'docx':
                content = self.extract_docx(file_path)
                return self.process_document(content, metadata, force_reindex)
            elif ext == 'txt':
                content = self.extract_txt(file_path)
                return self.process_document(content, metadata, force_reindex)
            elif ext in ['jpg', 'jpeg', 'png', 'webp']:
                # Les images sont toujours ré-indexées (pas de hash texte)
                image_info = self.extract_image(file_path)
                doc = Document(page_content=f"[IMAGE:{file_path}]", metadata={**metadata, **image_info})
                return [doc]
            else:
                logger.warning(f"Extension non supportée ignorée : .{ext} ({file_path})")
                return []
        except Exception as e:
            self.log_error(file_path, "ingest_file", str(e))
            return []

    async def ingest_url(
        self,
        url: str,
        metadata: Dict[str, Any],
        force_reindex: bool = False,
    ) -> List[Document]:
        content = await self.extract_webpage(url)
        return self.process_document(content, metadata, force_reindex)
        
    def ingest_db(self, db_uri: str, query: str, metadata: Dict[str, Any]) -> List[Document]:
        engine = create_engine(db_uri)
        with engine.connect() as conn:
            result = conn.execute(text(query)).fetchall()
            content = "\n".join([str(row) for row in result])
        return self.process_document(content, metadata)
